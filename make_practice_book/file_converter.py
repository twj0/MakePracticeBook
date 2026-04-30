"""
Input extraction pipeline for doc/docx/pdf/md/txt.
"""

from __future__ import annotations

import io
import os
import re
from pathlib import Path
from typing import Optional

from .doc_converter import convert_doc_to_docx


class FileConverter:
    """Convert supported input formats into normalized Markdown-like text."""

    def __init__(
        self,
        *,
        ocr_backend: str = "paddle",
        tesseract_cmd: Optional[str] = None,
        doc_strategy: str = "auto",
        temp_dir: Optional[Path] = None,
    ):
        self.ocr_backend = (ocr_backend or "paddle").strip().lower()
        self.doc_strategy = (doc_strategy or "auto").strip().lower()
        self.temp_dir = temp_dir or (Path.cwd() / "out" / "_tmp")
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        self._paddle_structure = None
        self._paddle_ocr = None
        self._trace: list[str] = []

        os.environ.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")

        if tesseract_cmd:
            import pytesseract  # type: ignore

            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    def convert_to_markdown(self, file_path: str) -> str:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        suffix = path.suffix.lower()
        if suffix == ".txt":
            self._log("Input detected: txt")
            return self._read_plain_text(path)
        if suffix == ".md":
            self._log("Input detected: markdown")
            return self._read_markdown(path)
        if suffix == ".doc":
            self._log(f"Input detected: legacy doc; converting via {self.doc_strategy}")
            converted = convert_doc_to_docx(path, self.temp_dir, strategy=self.doc_strategy)
            self._log(f"Converted doc to docx: {converted}")
            return self._convert_docx_to_markdown(converted)
        if suffix == ".docx":
            self._log("Input detected: docx")
            return self._convert_docx_to_markdown(path)
        if suffix == ".pdf":
            self._log(f"Input detected: pdf; OCR backend={self.ocr_backend}")
            return self._convert_pdf_to_markdown(path)

        raise ValueError("Unsupported file format. Supported: .doc, .docx, .pdf, .md, .txt")

    def _read_plain_text(self, path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="ignore")

    def _read_markdown(self, path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="ignore")

    def _convert_pdf_to_markdown(self, pdf_path: Path) -> str:
        if self.ocr_backend in {"paddle", "auto"}:
            paddle_text = self._convert_pdf_with_paddle(pdf_path)
            if paddle_text.strip():
                self._log("PDF extraction succeeded with Paddle path")
                return paddle_text

        import fitz  # type: ignore

        doc = fitz.open(pdf_path)
        markdown_content: list[str] = []
        try:
            for page_num, page in enumerate(doc, start=1):
                text = page.get_text("text")
                if self._should_ocr_text(text):
                    self._log(f"PDF page {page_num}: native text too sparse, falling back to page OCR")
                    text = self._ocr_page(page)
                if text.strip():
                    markdown_content.append(f"## Page {page_num}\n\n{text.strip()}\n")
        finally:
            doc.close()

        merged = "\n".join(markdown_content)
        if merged.strip():
            return merged

        raise RuntimeError("Failed to extract usable text from PDF")

    def _should_ocr_text(self, text: str) -> bool:
        stripped = text.strip()
        if not stripped:
            return True
        return len(stripped) < 80

    def _convert_pdf_with_paddle(self, pdf_path: Path) -> str:
        try:
            self._log("Trying PPStructureV3 layout parsing")
            pipeline = self._get_paddle_structure_pipeline()
            output = list(pipeline.predict(input=str(pdf_path)))
            structured = self._extract_markdown_from_ppstructure(output, pipeline)
            if structured.strip():
                self._log("PPStructureV3 produced structured markdown")
                return structured
            self._log("PPStructureV3 returned no usable markdown")
        except Exception as exc:
            self._log(f"PPStructureV3 failed: {exc}")

        try:
            self._log("Falling back to PaddleOCR page-by-page OCR")
            return self._convert_pdf_with_paddle_ocr(pdf_path)
        except Exception as exc:
            self._log(f"PaddleOCR page OCR failed: {exc}")
            return ""

    def _extract_markdown_from_ppstructure(self, output, pipeline) -> str:
        markdown_pages: list[object] = []
        text_pages: list[str] = []

        for res in output:
            for candidate in self._collect_result_objects(res):
                if isinstance(candidate, dict):
                    markdown_text = candidate.get("markdown")
                    if isinstance(markdown_text, str) and markdown_text.strip():
                        text_pages.append(markdown_text)
                        markdown_pages.append(candidate)
                    nested_res = candidate.get("res")
                    if isinstance(nested_res, dict):
                        nested_markdown = nested_res.get("markdown")
                        if isinstance(nested_markdown, str) and nested_markdown.strip():
                            text_pages.append(nested_markdown)
                elif isinstance(candidate, str) and candidate.strip():
                    text_pages.append(candidate)

        if markdown_pages and hasattr(pipeline, "concatenate_markdown_pages"):
            try:
                combined = pipeline.concatenate_markdown_pages(markdown_pages)
                if isinstance(combined, str) and combined.strip():
                    return combined
            except Exception:
                pass

        return "\n\n".join(item.strip() for item in text_pages if item.strip())

    def _convert_pdf_with_paddle_ocr(self, pdf_path: Path) -> str:
        import fitz  # type: ignore
        from PIL import Image  # type: ignore

        ocr = self._get_paddle_ocr_engine()

        doc = fitz.open(pdf_path)
        pages: list[str] = []
        try:
            for page_num, page in enumerate(doc, start=1):
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                image = Image.open(io.BytesIO(pix.tobytes("png")))
                page_text = self._run_paddle_ocr_on_image(image, ocr)
                if page_text.strip():
                    self._log(f"PaddleOCR extracted text from PDF page {page_num}")
                    pages.append(f"## Page {page_num}\n\n{page_text.strip()}\n")
        finally:
            doc.close()

        return "\n".join(pages)

    def _ocr_page(self, page) -> str:
        if self.ocr_backend == "none":
            return ""
        if self.ocr_backend == "paddle":
            try:
                from PIL import Image  # type: ignore
                import fitz  # type: ignore
            except ImportError:
                return ""

            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            image = Image.open(io.BytesIO(pix.tobytes("png")))
            ocr = self._get_paddle_ocr_engine()
            self._log("Using PaddleOCR for single-page fallback")
            return self._run_paddle_ocr_on_image(image, ocr)

        if self.ocr_backend != "tesseract" and self.ocr_backend != "auto":
            return ""

        try:
            import fitz  # type: ignore
            from PIL import Image  # type: ignore
            import cv2  # type: ignore
            import numpy as np  # type: ignore
            import pytesseract  # type: ignore
        except ImportError:
            return ""

        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        image = Image.open(io.BytesIO(pix.tobytes("png")))
        img_array = np.array(image)

        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array

        _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        denoised = cv2.fastNlMeansDenoising(binary)
        processed = Image.fromarray(denoised)
        return pytesseract.image_to_string(
            processed,
            lang="chi_sim+eng",
            config="--psm 6",
        )

    def _run_paddle_ocr_on_image(self, image, ocr) -> str:
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            temp_path = Path(tmp.name)
        try:
            image.save(temp_path)
            results = list(ocr.predict(str(temp_path)))
            lines: list[str] = []
            for result in results:
                for candidate in self._collect_result_objects(result):
                    if isinstance(candidate, dict):
                        source = candidate.get("res", candidate)
                        if isinstance(source, dict):
                            texts = source.get("rec_texts")
                            if isinstance(texts, list):
                                lines.extend(str(item).strip() for item in texts if str(item).strip())
                    elif isinstance(candidate, list):
                        lines.extend(str(item).strip() for item in candidate if str(item).strip())
            return "\n".join(lines)
        finally:
            temp_path.unlink(missing_ok=True)

    def _collect_result_objects(self, result) -> list[object]:
        candidates: list[object] = []
        for attr in ("markdown", "json", "res", "data"):
            if hasattr(result, attr):
                candidates.append(getattr(result, attr))
        if hasattr(result, "to_dict"):
            try:
                candidates.append(result.to_dict())
            except Exception:
                pass
        return candidates

    def _get_paddle_structure_pipeline(self):
        if self._paddle_structure is None:
            from paddleocr import PPStructureV3  # type: ignore

            self._paddle_structure = PPStructureV3(
                use_doc_orientation_classify=False,
                use_doc_unwarping=False,
                use_textline_orientation=False,
                use_table_recognition=False,
                use_formula_recognition=False,
                use_chart_recognition=False,
                use_seal_recognition=False,
                engine="paddle_dynamic",
                engine_config={"device_type": "cpu"},
            )
        return self._paddle_structure

    def _get_paddle_ocr_engine(self):
        if self._paddle_ocr is None:
            from paddleocr import PaddleOCR  # type: ignore

            self._paddle_ocr = PaddleOCR(
                use_doc_orientation_classify=False,
                use_doc_unwarping=False,
                use_textline_orientation=False,
                lang="ch",
                engine="paddle_dynamic",
                engine_config={"device_type": "cpu"},
            )
        return self._paddle_ocr

    def _convert_docx_to_markdown(self, docx_path: Path) -> str:
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:
            raise RuntimeError(
                "python-docx is required for DOCX conversion. Install it with `pip install python-docx`."
            ) from exc

        doc = Document(str(docx_path))
        self._log(f"Reading docx paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
        markdown_content: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style = (paragraph.style.name or "").lower()
            if "heading 1" in style or "title" in style:
                markdown_content.append(f"# {text}\n")
            elif "heading 2" in style:
                markdown_content.append(f"## {text}\n")
            elif "heading 3" in style:
                markdown_content.append(f"### {text}\n")
            elif "heading 4" in style:
                markdown_content.append(f"#### {text}\n")
            else:
                markdown_content.append(f"{text}\n")

        for table in doc.tables:
            markdown_content.append(self._convert_table_to_markdown(table))

        return "\n".join(markdown_content)

    def _convert_table_to_markdown(self, table) -> str:
        markdown_table = ["\n"]
        for index, row in enumerate(table.rows):
            cells = [self.clean_text(cell.text) for cell in row.cells]
            markdown_table.append("| " + " | ".join(cells) + " |")
            if index == 0:
                markdown_table.append("| " + " | ".join(["---"] * len(cells)) + " |")
        markdown_table.append("\n")
        return "\n".join(markdown_table)

    @staticmethod
    def clean_text(text: str) -> str:
        text = re.sub(r"\n\s*\n", "\n\n", text)
        text = re.sub(r" +", " ", text)
        return text.replace("\ufeff", "").strip()

    @property
    def trace(self) -> list[str]:
        return list(self._trace)

    def _log(self, message: str) -> None:
        self._trace.append(message)


def convert_file_to_markdown(
    input_path: str,
    *,
    ocr_backend: str = "paddle",
    tesseract_cmd: Optional[str] = None,
    doc_strategy: str = "auto",
) -> str:
    converter = FileConverter(
        ocr_backend=ocr_backend,
        tesseract_cmd=tesseract_cmd,
        doc_strategy=doc_strategy,
    )
    markdown_content = converter.convert_to_markdown(input_path)
    return FileConverter.clean_text(markdown_content)
